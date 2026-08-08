import streamlit as st
from google import genai
from google.genai import types
import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import io
import re
from PIL import Image
import cneb_primaria_datos as cneb

# ==============================================================================
# CONFIGURACIÓN DE LA PÁGINA STREAMLIT
# ==============================================================================
st.set_page_config(
    page_title="PlanificaPrimaria - Plataforma para Docentes de Aula",
    page_icon="🍎",
    layout="wide"
)

# ==============================================================================
# INYECCIÓN CSS/JS NIVEL INGENIERÍA PARA ELIMINACIÓN DE INSIGNIAS FLOTANTES
# ==============================================================================
st.markdown("""
<style>
    /* 1. ANULACIÓN Y COLAPSO ABSOLUTO DE ELEMENTOS DE STREAMLIT CLOUD */
    header, footer, [data-testid="stHeader"], [data-testid="stDecoration"], 
    [data-testid="stStatusWidget"], [data-testid="stViewerBadge"], 
    [data-testid="manage-app-button"], .stAppDeployButton, .viewerBadge_container__1613n,
    button[title*="Streamlit"], div[class*="stDeployButton"], div[class*="viewerBadge"], 
    div[class*="ViewerBadge"], a[class*="viewerBadge"], a[class*="ViewerBadge"], 
    div[class*="profile"], div[class*="Profile"], div[class*="crown"], div[class*="Crown"], 
    div[class*="hostBadge"], div[class*="HostBadge"], div[class*="badge"], div[class*="Badge"], 
    div[class*="floating"], div[class*="Floating"], a[href*="streamlit"], a[href*="share.streamlit.io"] {
        display: none !important;
        visibility: hidden !important;
        opacity: 0 !important;
        width: 0px !important;
        height: 0px !important;
        max-width: 0px !important;
        max-height: 0px !important;
        margin: 0 !important;
        padding: 0 !important;
        overflow: hidden !important;
        pointer-events: none !important;
        position: absolute !important;
        left: -9999px !important;
        top: -9999px !important;
        z-index: -9999 !important;
        transform: scale(0) !important;
    }
    
    /* 2. FONDO CLARO Y ELEGANTE PARA TODA LA PÁGINA Y PANTALLA DE LOGIN */
    html, body, [data-testid="stAppViewContainer"], .stApp {
        background-color: #F8FAFC !important;
        color: #0F172A !important;
    }
    
    .block-container {
        padding-top: 1.2rem !important;
        padding-bottom: 1rem !important;
        background-color: #F8FAFC !important;
    }
    
    /* Encabezados */
    .main-header {
        font-size: 2.3rem;
        color: #1E3A8A;
        text-align: center;
        font-weight: 800;
        margin-bottom: 0.2rem;
    }
    .sub-header {
        font-size: 1.05rem;
        color: #4B5563;
        text-align: center;
        margin-bottom: 1.5rem;
    }

    /* 3. CAMPOS DE ENTRADA Y TEXTOS LEGIBLES */
    .stTextInput input, .stTextArea textarea, .stSelectbox [data-baseweb="select"] {
        background-color: #FFFFFF !important;
        color: #0F172A !important;
        border: 1px solid #94A3B8 !important;
        border-radius: 8px !important;
    }
    .stTextInput label, .stTextArea label, .stSelectbox label, .stSlider label, p, span, h1, h2, h3, h4 {
        color: #0F172A !important;
        font-weight: 600 !important;
    }

    /* 4. COLORES EXCLUSIVOS Y VISIBLES PARA CADA BOTÓN DE HERRAMIENTA */
    
    /* Botón 1: PROYECTO DE APRENDIZAJE (VERDE) */
    div.st-key-btn_proyecto > button, button[key="btn_proyecto"] {
        background: linear-gradient(135deg, #10B981 0%, #059669 100%) !important;
        background-color: #059669 !important;
        border-radius: 12px !important;
        border: none !important;
        box-shadow: 0 4px 12px rgba(16, 185, 129, 0.4) !important;
    }
    div.st-key-btn_proyecto > button p, button[key="btn_proyecto"] p, div.st-key-btn_proyecto > button span {
        color: #FFFFFF !important;
        font-weight: 800 !important;
        font-size: 1.05rem !important;
    }

    /* Botón 2: UNIDAD SARA (PURPURA / MORADO) */
    div.st-key-btn_unidad > button, button[key="btn_unidad"] {
        background: linear-gradient(135deg, #8B5CF6 0%, #7C3AED 100%) !important;
        background-color: #7C3AED !important;
        border-radius: 12px !important;
        border: none !important;
        box-shadow: 0 4px 12px rgba(139, 92, 246, 0.4) !important;
    }
    div.st-key-btn_unidad > button p, button[key="btn_unidad"] p, div.st-key-btn_unidad > button span {
        color: #FFFFFF !important;
        font-weight: 800 !important;
        font-size: 1.05rem !important;
    }

    /* Botón 3: SESIÓN DE APRENDIZAJE (AZUL) */
    div.st-key-btn_sesion > button, button[key="btn_sesion"] {
        background: linear-gradient(135deg, #3B82F6 0%, #2563EB 100%) !important;
        background-color: #2563EB !important;
        border-radius: 12px !important;
        border: none !important;
        box-shadow: 0 4px 12px rgba(59, 130, 246, 0.4) !important;
    }
    div.st-key-btn_sesion > button p, button[key="btn_sesion"] p, div.st-key-btn_sesion > button span {
        color: #FFFFFF !important;
        font-weight: 800 !important;
        font-size: 1.05rem !important;
    }

    /* Botón 4: FICHA DE APLICACIÓN (NARANJA) */
    div.st-key-btn_ficha > button, button[key="btn_ficha"] {
        background: linear-gradient(135deg, #F97316 0%, #D97706 100%) !important;
        background-color: #D97706 !important;
        border-radius: 12px !important;
        border: none !important;
        box-shadow: 0 4px 12px rgba(249, 115, 22, 0.4) !important;
    }
    div.st-key-btn_ficha > button p, button[key="btn_ficha"] p, div.st-key-btn_ficha > button span {
        color: #FFFFFF !important;
        font-weight: 800 !important;
        font-size: 1.05rem !important;
    }

    /* Botón 5: AFICHE NANO BANANA (ROJO / CARMESÍ) */
    div.st-key-btn_afiche > button, button[key="btn_afiche"] {
        background: linear-gradient(135deg, #EF4444 0%, #DC2626 100%) !important;
        background-color: #DC2626 !important;
        border-radius: 12px !important;
        border: none !important;
        box-shadow: 0 4px 12px rgba(239, 68, 68, 0.4) !important;
    }
    div.st-key-btn_afiche > button p, button[key="btn_afiche"] p, div.st-key-btn_afiche > button span {
        color: #FFFFFF !important;
        font-weight: 800 !important;
        font-size: 1.05rem !important;
    }

    /* BOTÓN PRINCIPAL DE GENERACIÓN */
    div.stButton > button:not([key="btn_proyecto"]):not([key="btn_unidad"]):not([key="btn_sesion"]):not([key="btn_ficha"]):not([key="btn_afiche"]) {
        background: linear-gradient(135deg, #2563EB 0%, #1D4ED8 100%) !important;
        background-color: #2563EB !important;
        border-radius: 10px !important;
        border: none !important;
        box-shadow: 0 4px 12px rgba(37, 99, 235, 0.4) !important;
    }
    div.stButton > button:not([key="btn_proyecto"]):not([key="btn_unidad"]):not([key="btn_sesion"]):not([key="btn_ficha"]):not([key="btn_afiche"]) p {
        color: #FFFFFF !important;
        font-weight: 800 !important;
        font-size: 1.1rem !important;
    }
</style>
""", unsafe_allow_html=True)

# SCRIPT JAVASCRIPT GLOBAL QUE SE EJECUTA EN EL MARCO PADRE
st.markdown("""
<script>
function injectKillStyle() {
    const targets = [window.document, window.parent.document, window.top.document];
    targets.forEach(doc => {
        try {
            if (doc && !doc.getElementById('sys-kill-style')) {
                const style = doc.createElement('style');
                style.id = 'sys-kill-style';
                style.innerHTML = `
                    [data-testid="stViewerBadge"],
                    [data-testid="manage-app-button"],
                    .viewerBadge_container__1613n,
                    .stAppDeployButton,
                    div[class*="viewerBadge"],
                    div[class*="ViewerBadge"],
                    a[class*="viewerBadge"],
                    a[class*="ViewerBadge"],
                    div[class*="profile"],
                    div[class*="Profile"],
                    div[class*="crown"],
                    div[class*="Crown"],
                    div[class*="badge"],
                    div[class*="Badge"],
                    a[href*="streamlit"] {
                        display: none !important;
                        visibility: hidden !important;
                        opacity: 0 !important;
                        pointer-events: none !important;
                        width: 0 !important;
                        height: 0 !important;
                    }
                `;
                doc.head.appendChild(style);
            }
            const badSelectors = [
                '[data-testid="stViewerBadge"]',
                '[data-testid="manage-app-button"]',
                '.stAppDeployButton',
                '.viewerBadge_container__1613n',
                'a[href*="streamlit"]'
            ];
            badSelectors.forEach(s => {
                doc.querySelectorAll(s).forEach(el => el.remove());
            });
        } catch(e) {}
    });
}
setInterval(injectKillStyle, 150);
window.addEventListener('load', injectKillStyle);
</script>
""", unsafe_allow_html=True)

st.markdown('<div class="main-header">🍎 PlanificaPrimaria - Sistema para Docentes de Aula</div>', unsafe_allow_html=True)
st.markdown('<div class="sub-header">Plataforma Inteligente de Planificación Curricular para Educación Primaria (CNEB - MINEDU)</div>', unsafe_allow_html=True)

# ==============================================================================
# CONTROL DE ACCESO MEDIANTE CONTRASEÑA
# ==============================================================================
def check_password():
    """Retorna True si el usuario ingresó la contraseña correcta."""
    if "password_correct" not in st.session_state:
        st.session_state["password_correct"] = False

    if st.session_state["password_correct"]:
        return True

    st.markdown("---")
    col1, col2, col3 = st.columns([1, 2, 1])
    with col2:
        st.subheader("🔒 Acceso Restringido al Sistema")
        st.info("💡 Por favor, ingresa la contraseña para acceder a la plataforma.")
        pwd_input = st.text_input("Contraseña de acceso:", type="password", key="pwd_input")
        
        if st.button("Ingresar 🚀"):
            target_pwd = st.secrets.get("APP_PASSWORD", "docente2026")
            if pwd_input == target_pwd:
                st.session_state["password_correct"] = True
                st.rerun()
            else:
                st.error("❌ Contraseña incorrecta. Inténtalo de nuevo.")
    return False

if not check_password():
    st.stop()

# ==============================================================================
# INICIALIZACIÓN DE MEMORIA PERSISTENTE (st.session_state)
# ==============================================================================
if 'resultado_md' not in st.session_state:
    st.session_state['resultado_md'] = None
if 'tipo_doc_generado' not in st.session_state:
    st.session_state['tipo_doc_generado'] = None
if 'fname_clean' not in st.session_state:
    st.session_state['fname_clean'] = None
if 'ie_nombre_generado' not in st.session_state:
    st.session_state['ie_nombre_generado'] = None
if 'tipo_documento' not in st.session_state:
    st.session_state['tipo_documento'] = "Proyecto de Aprendizaje"
if 'imagen_nanobanana' not in st.session_state:
    st.session_state['imagen_nanobanana'] = None
if 'imagen_bytes' not in st.session_state:
    st.session_state['imagen_bytes'] = None

# ==============================================================================
# BARRA LATERAL (SIDEBAR) - CONFIGURACIÓN Y API KEY
# ==============================================================================
st.sidebar.title("⚙️ Configuración")

if st.sidebar.button("🔒 Cerrar Sesión"):
    st.session_state["password_correct"] = False
    st.rerun()

st.sidebar.markdown("---")

if "GEMINI_API_KEY" in st.secrets and st.secrets["GEMINI_API_KEY"]:
    api_key = st.secrets["GEMINI_API_KEY"]
    st.sidebar.success("🔑 API Key activada desde el servidor.")
else:
    api_key = st.sidebar.text_input(
        "🔑 Google AI Studio API Key:", 
        type="password", 
        help="Consigue tu clave gratuita en https://aistudio.google.com/app/apikey"
    )

model_choice = st.sidebar.selectbox(
    "Modelo de Gemini:", 
    ["gemini-3.6-flash", "gemini-3.5-flash", "gemini-2.5-flash", "gemini-2.0-flash"]
)

st.sidebar.markdown("---")
st.sidebar.info("""
**Alineamiento CNEB Perú:**
• RM N.° 649-2016-MINEDU
• Estándares y Desempeños CNEB Íntegros
• Turno Único: 2 a 3 Sesiones diarias de 90 min
• Nivel Educación Primaria (1.° a 6.° Grado)
• Tablas en Colores Pasteles Variados
""")

# ==============================================================================
# SELECCIÓN DE HERRAMIENTAS DE AULA EN LA PÁGINA PRINCIPAL
# ==============================================================================
st.markdown("### 📋 Selecciona la Herramienta de Aula que deseas elaborar:")

col_b1, col_b2, col_b3, col_b4, col_b5 = st.columns(5)

with col_b1:
    if st.button("🚀 Proyecto de Aprendizaje", key="btn_proyecto", use_container_width=True):
        st.session_state['tipo_documento'] = "Proyecto de Aprendizaje"
        st.rerun()

with col_b2:
    if st.button("📘 Unidad SARA", key="btn_unidad", use_container_width=True):
        st.session_state['tipo_documento'] = "Unidad de Aprendizaje (Modelo SARA)"
        st.rerun()

with col_b3:
    if st.button("🍎 Sesión de Aprendizaje", key="btn_sesion", use_container_width=True):
        st.session_state['tipo_documento'] = "Sesión de Aprendizaje"
        st.rerun()

with col_b4:
    if st.button("📝 Ficha de Aplicación", key="btn_ficha", use_container_width=True):
        st.session_state['tipo_documento'] = "Ficha de Aplicación / Trabajo (Para Alumnos)"
        st.rerun()

with col_b5:
    if st.button("🖼️ Afiche Nano Banana", key="btn_afiche", use_container_width=True):
        st.session_state['tipo_documento'] = "Afiche Educativo de la Sesión (Nano Banana)"
        st.rerun()

tipo_documento = st.session_state['tipo_documento']

COLOR_MAP = {
    "Proyecto de Aprendizaje": "#059669",
    "Unidad de Aprendizaje (Modelo SARA)": "#7C3AED",
    "Sesión de Aprendizaje": "#2563EB",
    "Ficha de Aplicación / Trabajo (Para Alumnos)": "#D97706",
    "Afiche Educativo de la Sesión (Nano Banana)": "#DC2626"
}
banner_color = COLOR_MAP.get(tipo_documento, "#059669")

st.markdown(f"""
<div style="background-color: {banner_color}; color: white; padding: 0.6rem 1rem; border-radius: 8px; font-weight: bold; font-size: 1.1rem; margin-top: 0.8rem; margin-bottom: 1.2rem; text-align: center;">
    📍 Herramienta Seleccionada: {tipo_documento.upper()}
</div>
""", unsafe_allow_html=True)

# ==============================================================================
# FUNCIONES AUXILIARES Y GENERADOR HÍBRIDO NANO BANANA
# ==============================================================================
def add_formatted_text(paragraph, text):
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.font.bold = True
        else:
            paragraph.add_run(part)

def markdown_to_docx(md_text, ie_nombre="I.E. N° 22303", es_horizontal=False):
    doc = docx.Document()
    PASTEL_COLORS = ['D9E1F2', 'E2EFDA', 'FFF2CC', 'E8D8F8', 'E0F2FE', 'FCE4D6']
    table_count = 0
    
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
        
        if es_horizontal:
            section.orientation = WD_ORIENT.LANDSCAPE
            section.page_width = Inches(11.69)
            section.page_height = Inches(8.27)
        else:
            section.orientation = WD_ORIENT.PORTRAIT
            section.page_width = Inches(8.27)
            section.page_height = Inches(11.69)
        
    p_box = doc.add_paragraph()
    p_box.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_box = p_box.add_run(f"🖼️ [ PEGAR AQUÍ LA INSIGNIA / ESCUDO DE LA {ie_nombre.upper()} ]\n")
    run_box.font.size = Pt(10)
    run_box.font.italic = True
    run_box.font.color.rgb = RGBColor(107, 114, 128)

    lines = md_text.split('\n')
    in_table = False
    table_data = []

    def render_table(t_data, color_hex):
        rows = len(t_data)
        cols = max(len(r) for r in t_data) if rows > 0 else 0
        if rows > 0 and cols > 0:
            t = doc.add_table(rows=rows, cols=cols)
            t.style = 'Table Grid'
            for r_idx, row_cells in enumerate(t_data):
                for c_idx, cell_value in enumerate(row_cells):
                    if c_idx < cols:
                        cell = t.cell(r_idx, c_idx)
                        p_cell = cell.paragraphs[0]
                        p_cell.text = ""
                        add_formatted_text(p_cell, cell_value)
                        
                        if r_idx == 0:
                            shading_elm = OxmlElement('w:shd')
                            shading_elm.set(qn('w:val'), 'clear')
                            shading_elm.set(qn('w:color'), 'auto')
                            shading_elm.set(qn('w:fill'), color_hex)
                            cell._tc.get_or_add_tcPr().append(shading_elm)
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.font.color.rgb = RGBColor(30, 58, 138)
                                    run.font.bold = True

    for line in lines:
        line_str = line.strip()
        line_str = re.sub(r'<br\s*/?>', ' ', line_str)
        line_str = re.sub(r'</?[a-zA-Z0-9]+\s*/>', ' ', line_str)
        line_str = re.sub(r'</?(table|tr|td|th|thead|tbody)[^>]*>', ' ', line_str, flags=re.IGNORECASE)
        
        if line_str.startswith('|') and line_str.endswith('|'):
            in_table = True
            if re.match(r'^\|[\s\:\-\|]+\|$', line_str):
                continue
            cells = [c.strip() for c in line_str.split('|')[1:-1]]
            table_data.append(cells)
            continue
        elif in_table:
            if table_data:
                table_count += 1
                header_color = PASTEL_COLORS[(table_count - 1) % len(PASTEL_COLORS)]
                render_table(table_data, header_color)
            in_table = False
            table_data = []

        heading_match = re.match(r'^(#{1,6})\s*(.*)$', line_str)
        if heading_match:
            hashes = heading_match.group(1)
            title_text = heading_match.group(2).strip()
            level = len(hashes)
            
            p = doc.add_paragraph()
            if level in [1, 2]:
                run = p.add_run(title_text.replace('**', ''))
                run.font.size = Pt(14)
                run.font.bold = True
                run.font.color.rgb = RGBColor(30, 58, 138)
            elif level in [3, 4]:
                run = p.add_run(title_text.replace('**', ''))
                run.font.size = Pt(12)
                run.font.bold = True
                run.font.color.rgb = RGBColor(30, 58, 138)
            else:
                add_formatted_text(p, title_text)
            continue

        if line_str.startswith('• ') or line_str.startswith('- '):
            p = doc.add_paragraph(style='List Bullet')
            clean_bullet = line_str[2:].strip()
            add_formatted_text(p, clean_bullet)
        elif line_str != "":
            p = doc.add_paragraph()
            add_formatted_text(p, line_str)

    if in_table and table_data:
        table_count += 1
        header_color = PASTEL_COLORS[(table_count - 1) % len(PASTEL_COLORS)]
        render_table(table_data, header_color)
        in_table = False
        table_data = []
            
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def generar_imagen_nanobanana(client, tema, grado, area):
    """Genera la imagen infográfica educativa con sistema dual (Imagen / Gemini Multimodal)"""
    prompt_nanobanana = f"""
    Full educational primary school session infographic poster (Estilo Sesión de Aprendizaje e Infografía Oficial MINEDU Perú).
    Grade: {grado}. Subject: {area}. Topic: '{tema}'.
    
    Visual Poster Layout & Structure:
    - TOP HEADER BANNER: Bold title "SESIÓN DE APRENDIZAJE: {tema.upper()}" with cute primary school children cartoon mascot icons.
    - TOP SECTION (DATOS Y PROPÓSITO): Small pastel information card boxes with checkmark icons detailing learning goals.
    - MAIN SECTION (DESARROLLO DE ACTIVIDADES EN SECUENCIA): Numbered activity step cards (1, 2, 3, 4) with timer clock icons, showing friendly Peruvian primary school students actively performing the learning activities step-by-step for '{tema}' in a safe classroom/school setting.
    - BOTTOM SECTION (REFLEXIÓN Y RECUERDA): A bottom "RECUERDA" bar with small safety tip badges and smile icons.
    
    Art Style: Highly detailed vector educational infographic poster layout, pastel blue/green/yellow/orange cards with rounded borders, white background, clean outlines, cute Peruvian primary school children illustrations, 3:4 vertical poster format.
    """
    
    ultimo_error = None
    
    # MÉTODO 1: Probar modelos de Imagen
    modelos_imagen = [
        'imagen-4.0-generate-001',
        'imagen-3.0-generate-002',
        'imagen-3.0-fast-generate-001',
        'imagen-3.0-generate-001'
    ]
    for mod in modelos_imagen:
        try:
            result = client.models.generate_images(
                model=mod,
                prompt=prompt_nanobanana,
                config=types.GenerateImagesConfig(
                    number_of_images=1,
                    output_mime_type="image/jpeg",
                    aspect_ratio="3:4",
                )
            )
            if hasattr(result, 'generated_images') and result.generated_images:
                for gen_img in result.generated_images:
                    img_bytes = gen_img.image.image_bytes
                    img = Image.open(io.BytesIO(img_bytes))
                    return img, img_bytes, None
        except Exception as err:
            ultimo_error = str(err)
            continue

    # MÉTODO 2: Probar generación multimodal nativa vía generate_content
    modelos_multimodal = ['gemini-2.0-flash-exp', 'gemini-2.0-flash']
    for mod in modelos_multimodal:
        try:
            response = client.models.generate_content(
                model=mod,
                contents=f"Genera la imagen de un afiche educativo visual: {prompt_nanobanana}",
                config=types.GenerateContentConfig(
                    response_modalities=["TEXT", "IMAGE"]
                )
            )
            if hasattr(response, 'candidates') and response.candidates:
                for part in response.candidates[0].content.parts:
                    if hasattr(part, 'inline_data') and part.inline_data and part.inline_data.data:
                        img_bytes = part.inline_data.data
                        img = Image.open(io.BytesIO(img_bytes))
                        return img, img_bytes, None
        except Exception as err:
            ultimo_error = str(err)
            continue

    return None, None, ultimo_error

# ==============================================================================
# FORMULARIO DE DATOS DE AULA
# ==============================================================================
st.subheader(f"📝 Configuración de Datos: {tipo_documento}")

c1, c2, c3 = st.columns(3)
with c1:
    dre_ugel = st.text_input("DRE / UGEL:", "Ica / Ica")
    ie_nombre = st.text_input("Institución Educativa:", "N° 22303 'Santa Rosa de Lima'")
with c2:
    director = st.text_input("Director:", "Lic. Bernardo Francisco Salcedo Barrientos")
    subdirector = st.text_input("Subdirector(es):", "Mg. Mariela Velásquez Cárdenas / Mg. Frank Bernaola Pérez")
with c3:
    docente = st.text_input("Docente de Aula:", "Sara María Quiroz Rodríguez")
    grado_seccion = st.selectbox("Grado y Sección:", ["1er Grado A", "2do Grado A", "3er Grado A", "4to Grado A", "5to Grado A", "6to Grado A"], index=2)

if tipo_documento in ["Sesión de Aprendizaje", "Ficha de Aplicación / Trabajo (Para Alumnos)", "Afiche Educativo de la Sesión (Nano Banana)"]:
    f1, f2, f3, f4 = st.columns(4)
    with f1:
        num_doc = st.text_input("N.° de Documento / Sesión / Ficha / Afiche:", "01")
    with f2:
        area_sel = st.selectbox("Área Curricular:", cneb.obtener_lista_areas(), index=0)
    with f3:
        fecha_sugerida = st.text_input("Fecha:", "05 de mayo de 2026")
    with f4:
        duracion_sesion = st.selectbox("Duración de la Sesión:", ["45 minutos", "90 minutos", "135 minutos"], index=1)
    
    fechas_duracion = fecha_sugerida
    duracion_semanas = 1

elif tipo_documento == "Proyecto de Aprendizaje":
    f1, f2, f3 = st.columns(3)
    with f1:
        num_doc = st.text_input("N.° de Proyecto:", "01")
    with f2:
        fechas_duracion = st.text_input("Fechas / Duración:", "Del 11 de marzo al 12 de abril de 2026 (4 Semanas)")
    with f3:
        duracion_semanas = st.slider("Número de Semanas del Proyecto:", min_value=2, max_value=5, value=4)
        area_sel = "Multidisciplinar"
        duracion_sesion = "90 minutos"

else:  # Unidad SARA
    f1, f2, f3 = st.columns(3)
    with f1:
        num_doc = st.text_input("N.° de Unidad:", "01")
    with f2:
        fechas_duracion = st.text_input("Fechas / Duración:", "Del 01 de abril al 03 de mayo de 2026 (5 Semanas)")
    with f3:
        duracion_semanas = st.slider("Número de Semanas de la Unidad:", min_value=2, max_value=5, value=5)
        area_sel = "Multidisciplinar"
        duracion_sesion = "90 minutos"

if tipo_documento in ["Sesión de Aprendizaje", "Ficha de Aplicación / Trabajo (Para Alumnos)", "Afiche Educativo de la Sesión (Nano Banana)"]:
    problema_contexto = st.text_input(
        "📌 Tema / Título de la Actividad, Ficha de Trabajo o Afiche:",
        value="Mis derechos y deberes"
    )
    titulo_opcional = ""
else:
    problema_contexto = st.text_area(
        "🚨 Problema, Situación Significativa o Actividades Propuestas (Puedes colocar tu Situación Significativa / Actividades completas para que la IA las respete, o escribir solo el problema de contexto para que la IA las genere automáticamente):",
        height=130,
        value="Poco hábito de recolección de residuos sólidos y acumulación de botellas de plástico en el patio durante el recreo por parte de los estudiantes de 3er grado."
    )
    titulo_opcional = st.text_input("Título Opcional (Déjalo en blanco si deseas que la IA cree un título creativo automático):", value="")

# ==============================================================================
# PROMPTS MAESTROS COMPLETOS ALINEADOS AL CNEB
# ==============================================================================
def generar_prompt_sesion():
    if "45" in duracion_sesion:
        t_inicio, t_desarrollo, t_cierre = "10 min", "30 min", "5 min"
    elif "135" in duracion_sesion:
        t_inicio, t_desarrollo, t_cierre = "20 min", "100 min", "15 min"
    else:
        t_inicio, t_desarrollo, t_cierre = "20 min", "60 min", "10 min"

    return f"""
Actúa como: Un Especialista en Currículo Nacional de Educación Básica (CNEB) del MINEDU, experto en planificación pedagógica de nivel Primaria.
Tu objetivo: Elaborar una sesión de aprendizaje completa siguiendo estrictamente el formato y estructura del modelo proporcionado.

Datos para la sesión (Configuración):
• Grado y Sección: {grado_seccion}
• Área Curricular: {area_sel}
• Tema/Título de la sesión: {problema_contexto}
• Fecha sugerida: {fecha_sugerida}
• DRE / UGEL: {dre_ugel}
• Institución Educativa: {ie_nombre}
• Director: {director}
• Subdirector(es): {subdirector}
• Docente de Aula: {docente}
• Duración Total: {duracion_sesion}

INSTRUCCIONES DE FORMATO Y CONTENIDO (OBLIGATORIO):
• Estructura de Cuadros: Utiliza exactamente los mismos cuadros del modelo (Datos Informativos, Propósitos, Enfoques, Metas, Preparación, Momentos de la sesión y Escala de Valoración). NO INCLUYAS NINGUNA SITUACIÓN SIGNIFICATIVA.
• Alineación CNEB: Selecciona la Competencia, Capacidades y Desempeños (precisados si es necesario) directamente del Programa Curricular de Educación Primaria del MINEDU correspondiente al grado ({grado_seccion}). Para el Estándar de Aprendizaje del CNEB, escríbelo EN SU TOTALIDAD Y DE MANERA ÍNTEGRA sin ningún corte, resumen ni omisión, resaltando en **negrita** únicamente el fragmento trabajado.
• Criterios de Evaluación: Deben redactarse bajo la estructura implícita de ACCIÓN + CONTENIDO + CONDICIÓN.
• Redacción de Actividades: Las actividades en los momentos de Inicio, Desarrollo y Cierre deben estar redactadas en PRIMERA PERSONA DEL PLURAL Y TIEMPO PRESENTE.

ESTRUCTURA DE SALIDA REQUERIDA:
# **SESIÓN DE APRENDIZAJE N.º {num_doc}**
## **{problema_contexto.upper()}**

• TABLA I: DATOS INFORMATIVOS (ESTRICTAMENTE EN 2 COLUMNAS)
| DATOS INFORMATIVOS | DETALLE / INFORMACIÓN |
| DRE / UGEL | {dre_ugel} |
| Institución Educativa | {ie_nombre} |
| Director | {director} |
| Subdirector(es) | {subdirector} |
| Docente de Aula | {docente} |
| Grado y Sección | {grado_seccion} |
| Área Curricular | {area_sel} |
| Fecha | {fecha_sugerida} |
| Duración | {duracion_sesion} |

• TABLA II: PROPÓSITOS DE APRENDIZAJE Y EVIDENCIAS
| ÁREA | COMPETENCIA Y CAPACIDADES | ESTÁNDAR DE APRENDIZAJE (CNEB completo con **negrita**) | DESEMPEÑOS PRECISADOS (CNEB) | CRITERIOS DE EVALUACIÓN | PROPÓSITO DE LA SESIÓN | EVIDENCIA DE APRENDIZAJE | INSTRUMENTO DE EVALUACIÓN |

• TABLA III: ENFOQUES TRANSVERSALES
| ENFOQUE TRANSVERSAL | VALORES | ACTITUDES OBSERVABLES |

• TABLA IV: COMPETENCIA TRANSVERSAL
| COMPETENCIA TRANSVERSAL | CAPACIDADES | DESEMPEÑOS PRECISADOS |

• TABLA V: META DE APRENDIZAJE
| META DE APRENDIZAJE ({grado_seccion}) | DESCRIPCIÓN DE LA META |

• TABLA VI: PREPARACIÓN DE LA SESIÓN
| ¿Qué necesitamos hacer antes de la sesión? | ¿Qué recursos o materiales se utilizarán en esta sesión? |

• MOMENTOS DE LA SESIÓN:
- **INICIO ({t_inicio})**
- **DESARROLLO ({t_desarrollo})**
- **CIERRE ({t_cierre})**

• TABLA VII: ESCALA DE VALORACIÓN (30 alumnos ficticios)
"""

def generar_prompt_ficha_trabajo():
    return f"""
Actúa como: Especialista en Educación Primaria (CNEB - MINEDU Perú) y diseñador experto de material educativo impreso.
Elabora una FICHA DE TRABAJO / APLICACIÓN PARA EL ESTUDIANTE sobre {problema_contexto} para {grado_seccion} en el área de {area_sel}.

ESTRUCTURA DE SALIDA REQUERIDA (MARKDOWN PURA EN TABLAS):
# **FICHA DE TRABAJO DE {area_sel.upper()} N.º {num_doc}**
## **{problema_contexto.upper()}**

## **DATOS INFORMATIVOS**
| DATOS INFORMATIVOS | DETALLE / INFORMACIÓN |
| Institución Educativa | {ie_nombre} |
| Grado y Sección | {grado_seccion} |
| Área Curricular | {area_sel} |
| Docente de Aula | {docente} |
| Fecha | {fecha_sugerida} |
| Estudiante | __________________________________________________ |

## **PROPÓSITO DE HOY**
[Propósito amigable para el estudiante]

• SECCIÓN 1: "ME PREPARO Y DESCUBRO"
• SECCIÓN 2: "MANOS A LA OBRA / APLICO LO APRENDIDO"
• SECCIÓN 3: "MI RETO FINAL / MI COMPROMISO"

• TABLA II: AUTOEVALUACIÓN DE MIS LOGROS
"""

def generar_prompt_proyecto():
    val_titulo = f'"{titulo_opcional}"' if titulo_opcional.strip() else 'Crea un TÍTULO innovador y creativo para el proyecto basado en el problema.'

    return f"""
Actúa como un docente especialista de Primaria MINEDU Perú. Elabora un PROYECTO DE APRENDIZAJE completo.
PROHIBIDO usar símbolos #### o ##### y etiquetas HTML. Usa Markdown limpio y estructura strictly en TABLAS Y CUADROS.

A PARTIR DEL PROBLEMA DEL CONTEXTO DEL DOCENTE:
{problema_contexto}

OBLIGATORIO - GENERACIÓN AUTOMÁTICA DE TÍTULO Y SITUACIÓN SIGNIFICATIVA:
1. Genera un TÍTULO del proyecto: {val_titulo}
2. Redacta la SITUACIÓN SIGNIFICATIVA COMPLETA estructurada en 3 párrafos.

ORDEN ESTRUCTURAL ESTRICTO DE SALIDA (Sigue exactamente esta secuencia):

1. ENCABEZADO Y TABLA I: DATOS INFORMATIVOS (Muestra exactamente: DRE/UGEL: {dre_ugel}, IE: {ie_nombre}, Director: {director}, Subdirector: {subdirector}, Docente: {docente}, Grado/Sección: {grado_seccion}, Duración: {fechas_duracion}).

2. SITUACIÓN SIGNIFICATIVA GENERADA (Ubicada OBLIGATORIAMENTE justo debajo de los Datos Informativos).

3. PLANIFICACIÓN DEL PROYECTO CON LOS ESTUDIANTES (Tabla: ¿Qué haremos?, ¿Qué sabemos?, ¿Qué queremos saber?, ¿Cómo lo haremos?, ¿Qué necesitamos?, ¿Cómo nos organizamos?).

4. MATRIZ DE PROPÓSITOS DE APRENDIZAJE POR SEMANA (SECCIÓN CONTINUA COMPLETA DESDE LA SEMANA 1 HASTA LA SEMANA {duracion_semanas}):
   - Para CADA una de las {duracion_semanas} semanas, coloca el **TÍTULO CREATIVO DE LA ACTIVIDAD DE LA SEMANA** (Ejemplo: SEMANA 1: "Investigamos los problemas ambientales de nuestro colegio").
   - Debajo del título de la semana, presenta la Matriz de Propósitos en sus 8 COLUMNAS EXACTAS:
     | ÁREA | ACTIVIDAD | COMPETENCIA Y CAPACIDADES | ESTÁNDAR DE APRENDIZAJE | DESEMPEÑO PRECISADO | CRITERIOS DE EVALUACIÓN | EVIDENCIA | INSTRUMENTO DE EVALUACIÓN |
   - REGLA OBLIGATORIA DEL ESTÁNDAR: Copia el **ESTÁNDAR DE APRENDIZAJE EN SU TOTALIDAD Y DE MANERA ÍNTEGRA** tal cual figura en el CNEB oficial (RM N.º 649-2016-MINEDU) sin ningún corte ni resumen, y RESALTA EN **NEGRITA** (`**la parte específica movilizada en la actividad**`).
   - Copia el DESEMPEÑO ÍNTEGRO del CNEB con la parte trabajada en **negrita**.
   - REGLA OBLIGATORIA DE COBERTURA DE ÁREAS EN LA MATRIZ: En cada una de las {duracion_semanas} semanas, debes incluir OBLIGATORIAMENTE filas para TODAS Y CADA UNA DE LAS ÁREAS CURRICULARES SIN EXCEPCIÓN: Comunicación (3 comp.), Matemática (4 comp.), Personal Social, Ciencia y Tecnología, Educación Religiosa, Arte y Cultura, Educación Física y Tutoría / Competencias Transversales.

5. SECUENCIA DE ACTIVIDADES CON LOS DÍAS COMO COLUMNAS DE TABLA (2 A 3 SESIONES DIARIAS DE 90 MINUTOS EN TURNO ÚNICO - 10 A 15 SESIONES POR SEMANA):
   - Presenta esta sección OBLIGATORIAMENTE AL TÉRMINO DE TODA LA MATRIZ DE PROPÓSITOS.
   - Para cada semana (Semana 1 a {duracion_semanas}), coloca el **TÍTULO DE LA SEMANA** y crea una TABLA OBLIGATORIA donde LAS COLUMNAS SEAN LOS DÍAS DE LA SEMANA:
     | LUNES | MARTES | MIÉRCOLES | JUEVES | VIERNES |
   - En cada casillero diario, programa de 2 a 3 sesiones de 90 minutos en turno único (sin dividir en mañana/tarde), indicando el **ÁREA CURRICULAR DESTACADA**:
     • Sesión 1 (90 min): **[ÁREA]**: [Competencia específica] - [Actividad en 1ª persona plural]
     • Sesión 2 (90 min): **[ÁREA]**: [Competencia específica] - [Actividad en 1ª persona plural]
     • Sesión 3 (90 min, si aplica): **[ÁREA]**: [Competencia específica] - [Actividad en 1ª persona plural]
   - REGLA OBLIGATORIA DE ÁREAS EN LA SECUENCIA DE ACTIVIDADES: En la tabla semanal de actividades, DEBES DISTRIBUIR Y CONSIDERAR OBLIGATORIAMENTE TODAS Y CADA UNA DE LAS ÁREAS CURRICULARES EN CADA SEMANA SIN EXCEPCIÓN (Comunicación, Matemática, Personal Social, Ciencia y Tecnología, Religión, Arte, Educación Física y Tutoría).

6. TABLA DE ENFOQUES TRANSVERSALES.
7. PRODUCTO FINAL TANGIBLE DEL PROYECTO.
8. LISTA CLASIFICADA DE MATERIALES Y RECURSOS.
9. TABLA VIII: REFLEXIONES SOBRE LOS APRENDIZAJES (Tabla final obligatoria).
"""

def generar_prompt_unidad_sara():
    val_titulo = f'"{titulo_opcional}"' if titulo_opcional.strip() else 'Crea un TÍTULO motivador para la Unidad de Aprendizaje basado en el contexto/problema.'

    return f"""
Actúa como docente especialista de Primaria MINEDU Perú. Elabora una UNIDAD DE APRENDIZAJE completa y detallada (Modelo SARA).
PROHIBIDO usar símbolos #### o ##### y etiquetas HTML. Usa Markdown limpio y estructura strictly en TABLAS Y CUADROS.

ENTRADA PROVISTA POR EL DOCENTE (PROBLEMA DE CONTEXTO, SITUACIÓN SIGNIFICATIVA COMPLETA Y/O ACTIVIDADES PROPUESTAS):
{problema_contexto}

REGLA DE PROCESAMIENTO DE LA SITUACIÓN SIGNIFICATIVA Y ACTIVIDADES:
1. SI EL DOCENTE INGRESÓ UNA SITUACIÓN SIGNIFICATIVA COMPLETA O ACTIVIDADES ESPECÍFICAS: Utiliza, respeta y adapta fielmente dicho texto e ideas dentro de la sección "II. SITUACIÓN (RETO)" y en la matriz curricular.
2. SI EL DOCENTE SOLO INGRESÓ UN PROBLEMA O INTERÉS DEL CONTEXTO BREVE: La IA debe GENERAR AUTOMÁTICAMENTE la Situación Significativa completa estructurada en 3 párrafos (Párrafo 1: Diagnóstico de la problemática local/nacional; Párrafo 2: Propuesta pedagógica e integración de áreas; Párrafo 3: Interrogantes retadoras y desafíos) y articular las actividades sugeridas correspondientes.

ORDEN ESTRUCTURAL ESTRICTO DE SALIDA (Sigue exactamente esta secuencia):

UNIDAD DE APRENDIZAJE N.º {num_doc}
{val_titulo}

I. TABLA I: DATOS GENERALES (Muestra exactamente: DRE/UGEL: {dre_ugel}, IE: {ie_nombre}, Director: {director}, Subdirector: {subdirector}, Docente: {docente}, Grado/Sección: {grado_seccion}, Fechas y Duración: {fechas_duracion}, Duración en Semanas: {duracion_semanas} semanas).

II. SITUACIÓN (RETO):
(Muestra la SITUACIÓN SIGNIFICATIVA provista por el docente o la generada automáticamente por la IA con sus 3 párrafos y retos justo debajo de los Datos Generales).

III. ENFOQUES TRANSVERSALES:
| ENFOQUE TRANSVERSAL | VALOR | ACTITUD |

IV. ACTIVIDADES PERTINENTES AL PROPÓSITO DE APRENDIZAJE:
(Presenta una lista ordenada de las grandes actividades pedagógicas planificadas por semana, respetando las propuestas por el docente si las incluyó).

V. PROPÓSITO DE APRENDIZAJE, CRITERIOS DE EVALUACIÓN Y ACTIVIDADES SUGERIDAS (MATRIZ DE APRENDIZAJES POR ÁREA):
   - Presenta la Matriz Curricular completa en sus 8 COLUMNAS EXACTAS:
     | ÁREA | ACTIVIDAD | COMPETENCIA Y CAPACIDADES | ESTÁNDAR DE APRENDIZAJE | DESEMPEÑO PRECISADO | CRITERIOS DE EVALUACIÓN | EVIDENCIA | INSTRUMENTO DE EVALUACIÓN |
   
   - REGLA CRÍTICA OBLIGATORIA PARA MATEMÁTICA Y COMUNICACIÓN:
     1. En el área de MATEMÁTICA debes abordar e incluir OBLIGATORIAMENTE LAS 4 COMPETENCIAS del CNEB distribuida a lo largo de la unidad:
        - Resuelve problemas de cantidad.
        - Resuelve problemas de regularidad, equivalencia y cambio.
        - Resuelve problemas de forma, movimiento y localización.
        - Resuelve problemas de gestión de datos e incertidumbre.
     2. En el área de COMUNICACIÓN debes abordar e incluir OBLIGATORIAMENTE LAS 3 COMPETENCIAS del CNEB distribuida a lo largo de la unidad:
        - Se comunica oralmente en su lengua materna.
        - Lee diversos tipos de textos escritos en su lengua materna.
        - Escribe diversos tipos de textos en su lengua materna.

   - REGLA STRICTA Y ABSOLUTA PARA EL ESTÁNDAR DE APRENDIZAJE:
     1. El Estándar de Aprendizaje del ciclo correspondiente debe escribirse TAL CUAL figura de forma oficial en el CNEB (RM N.º 649-2016-MINEDU).
     2. Queda STRICTAMENTE PROHIBIDO modificar, parafrasear, resumir, cortar u omitir cualquier parte del texto del estándar. Debe incluirse el texto completo e íntegro del estándar del ciclo.
     3. ÚNICAMENTE debes resaltar en NEGRITA (`**texto en negrita**`) el fragmento o porción específica del estándar que se está abordando o movilizando en esa actividad. El resto del texto del estándar debe permanecer en texto plano normal.

   - REGLA STRICTA Y ABSOLUTA PARA EL DESEMPEÑO:
     1. El Desempeño correspondiente al grado debe escribirse TAL CUAL figura de forma oficial en el Programa Curricular de Educación Primaria del CNEB.
     2. Queda STRICTAMENTE PROHIBIDO resumir, omitir o recortar el texto del desempeño. Debe redactarse de manera completa e íntegra.
     3. ÚNICAMENTE debes resaltar en NEGRITA (`**texto precisado**`) la precisión del desempeño o la porción específica que se está trabajando directamente en la actividad. El resto del desempeño debe permanecer en texto plano normal.
   
   - Formula CRITERIOS DE EVALUACIÓN claros basados en: Verbo de acción + Contenido disciplinar + Condición/Contexto.
   - REGLA OBLIGATORIA DE COBERTURA DE ÁREAS EN LA MATRIZ: En cada una de las {duracion_semanas} semanas, debes incluir OBLIGATORIAMENTE filas para TODAS Y CADA UNA DE LAS ÁREAS CURRICULARES SIN EXCEPCIÓN: Comunicación (3 competencias), Matemática (4 competencias), Personal Social, Ciencia y Tecnología, Educación Religiosa, Arte y Cultura, Educación Física y Tutoría / Competencias Transversales.

VI. TUTORÍA Y EDUCACIÓN EDUCATIVA:
| DIMENSIÓN | SESIÓN | ¿QUÉ BUSCAMOS? |

VII. COMPETENCIAS TRANSVERSALES:
- "Se desenvuelve en los entornos virtuales generados por las TIC" (capacidades y desempeños precisados).
- "Gestiona su aprendizaje de manera autónoma" (capacidades y desempeños precisados).

VIII. PROGRAMACIÓN DE ACTIVIDADES / SECUENCIA CRONOLÓGICA DE ACTIVIDADES SUGERIDAS (SEMANA A SEMANA):
   - Presenta esta sección OBLIGATORIAMENTE AL TÉRMINO DE TODA LA MATRIZ DE APRENDIZAJES.
   - Para cada semana (Semana 1 a {duracion_semanas}), coloca el **TÍTULO DE LA SEMANA** y crea una TABLA OBLIGATORIA donde LAS COLUMNAS SEAN LOS DÍAS DE LA SEMANA:
     | LUNES | MARTES | MIÉRCOLES | JUEVES | VIERNES |
   - En cada casillero diario, programa de 2 a 3 sesiones de 90 minutos en turno único (sin dividir en mañana/tarde), indicando el **ÁREA CURRICULAR DESTACADA**:
     • Sesión 1 (90 min): **[ÁREA]**: [Competencia específica] - [Actividad en 1ª persona plural]
     • Sesión 2 (90 min): **[ÁREA]**: [Competencia específica] - [Actividad en 1ª persona plural]
     • Sesión 3 (90 min, si aplica): **[ÁREA]**: [Competencia específica] - [Actividad en 1ª persona plural]
   - REGLA OBLIGATORIA DE ÁREAS EN LA SECUENCIA DE ACTIVIDADES: En la tabla semanal de actividades, DEBES DISTRIBUIR Y CONSIDERAR OBLIGATORIAMENTE TODAS Y CADA UNA DE LAS ÁREAS CURRICULARES EN CADA SEMANA SIN EXCEPCIÓN.

IX. MATERIALES BÁSICOS Y RECURSOS A UTILIZAR:
- Para el estudiante.
- Para el docente.

X. REFLEXIONES SOBRE LOS APRENDIZAJES:
- Incluye la tabla o lista de preguntas de reflexión y metacognición del docente sobre el desarrollo de la unidad.
"""

# ==============================================================================
# EJECUCIÓN CON GOOGLE AI STUDIO (GEMINI API Y NANO BANANA)
# ==============================================================================
st.markdown("---")

if st.button(f"✨ Generar {tipo_documento}"):
    if not api_key:
        st.error("⚠️ Ingresa tu API Key de Google AI Studio en la barra lateral izquierda o en los Secrets.")
    elif not problema_contexto:
        st.warning("⚠️ Completa el campo del Tema, Problema o Situación Significativa.")
    else:
        try:
            client = genai.Client(api_key=api_key)
            
            # SI SE SELECCIONA EL AFICHE DE NANO BANANA:
            if tipo_documento == "Afiche Educativo de la Sesión (Nano Banana)":
                st.session_state['tipo_doc_generado'] = tipo_documento
                with st.spinner("🎨 Nano Banana está diseñando la Lámina / Afiche Educativo Ilustrado en HD para tu sesión..."):
                    img_obj, img_bytes, err_detallado = generar_imagen_nanobanana(
                        client, 
                        tema=problema_contexto, 
                        grado=grado_seccion, 
                        area=area_sel
                    )
                    
                    if img_obj is not None:
                        st.session_state['imagen_nanobanana'] = img_obj
                        st.session_state['imagen_bytes'] = img_bytes
                        st.session_state['resultado_md'] = f"""
# 🖼️ **AFICHE EDUCATIVO DE LA SESIÓN (NANO BANANA AI)**
**Tema:** {problema_contexto} | **Área:** {area_sel} | **Grado:** {grado_seccion}  
**Institución Educativa:** {ie_nombre} | **Docente:** {docente}  

---
*El afiche ilustrado ha sido generado en alta resolución. Puedes observarlo en la vista previa y descargarlo directamente en formato JPG listo para imprimir o proyectar en el aula.*
"""
                        st.success("✅ ¡Afiche Educativo Ilustrado generado con éxito!")
                    else:
                        st.session_state['imagen_nanobanana'] = None
                        st.session_state['imagen_bytes'] = None
                        st.session_state['resultado_md'] = f"⚠️ No se pudo generar la imagen del afiche debido a una restricción de la API de Google AI Studio.\n\n**Detalle técnico:** {err_detallado}"
                        st.error(f"❌ Ocurrió un problema de la API de Google AI Studio al generar la imagen. Detalle técnico: {err_detallado}")

            # SI SE SELECCIONA OTRA HERRAMIENTA (PROYECTO, UNIDAD, SESIÓN, FICHA):
            else:
                if tipo_documento == "Sesión de Aprendizaje":
                    prompt_maestro = generar_prompt_sesion()
                    sys_inst = "Eres un Especialista Curricular de Educación Primaria del MINEDU Perú. Creas sesiones de aprendizaje en tablas sin incluir situación significativa, incluyendo datos informativos en 2 columnas, propósitos de aprendizaje, enfoques, competencia transversal, meta de aprendizaje, preparación, momentos con procesos didácticos del área en 1ra persona plural tiempo presente, y escala de valoración con 30 estudiantes ficticios."
                elif tipo_documento == "Ficha de Aplicación / Trabajo (Para Alumnos)":
                    prompt_maestro = generar_prompt_ficha_trabajo()
                    sys_inst = "Eres un Especialista Curricular y Diseñador de Material Educativo de Educación Primaria del MINEDU Perú. Creas fichas de trabajo aplicando el proceso didáctico del área elegida. Muestras 'DATOS INFORMATIVOS' y 'PROPÓSITO DE HOY' obligatoriamente como SUBTÍTULOS FUERA DE LAS TABLAS. PROHIBIDO USAR ETIQUETAS HTML COMO <tr>, <td>, <th>, <table>, <tbody>."
                elif tipo_documento == "Proyecto de Aprendizaje":
                    prompt_maestro = generar_prompt_proyecto()
                    sys_inst = "Eres un Especialista Curricular de Educación Primaria del MINEDU Perú."
                else:
                    prompt_maestro = generar_prompt_unidad_sara()
                    sys_inst = (
                        "Eres un Especialista Curricular de Educación Primaria del MINEDU Perú. "
                        "Elaboras Unidades de Aprendizaje completas en formato Markdown. "
                        "REGLA CRÍTICA PARA MATEMÁTICA Y COMUNICACIÓN: Debes incluir OBLIGATORIAMENTE las 4 competencias del área de Matemática y las 3 competencias del área de Comunicación a lo largo de la unidad. "
                        "REGLA CRÍTICA PARA EL ESTÁNDAR Y DESEMPEÑO: Debes copiar el texto completo e íntegro tanto del Estándar de Aprendizaje como del Desempeño oficial del CNEB (RM N.° 649-2016-MINEDU) para el grado/ciclo, sin modificar, resumir, alterar ni recortar ninguna palabra. "
                        "Resalta en NEGRITA (**texto**) únicamente el fragmento o precisión que se moviliza o evalúa en la actividad. El resto del texto del estándar y del desempeño debe permanecer exactamente en texto normal. "
                        "Si el docente proporciona su propia Situación Significativa o actividades, utilízalas y respétalas íntegramente; si solo indica un problema breve, genera la Situación Significativa automáticamente."
                    )
                    
                with st.spinner(f"🧠 Google Gemini ({model_choice}) está procesando tu {tipo_documento} para {grado_seccion}..."):
                    config = types.GenerateContentConfig(
                        system_instruction=sys_inst,
                        temperature=0.2
                    )
                    try:
                        response = client.models.generate_content(
                            model=model_choice,
                            contents=prompt_maestro,
                            config=config
                        )
                    except Exception as model_err:
                        err_text = str(model_err)
                        if "404" in err_text or "NOT_FOUND" in err_text:
                            response = client.models.generate_content(
                                model="gemini-2.0-flash",
                                contents=prompt_maestro,
                                config=config
                            )
                        else:
                            raise model_err
                    
                    st.session_state['resultado_md'] = response.text
                    st.session_state['tipo_doc_generado'] = tipo_documento
                    st.session_state['fname_clean'] = f"{tipo_documento.replace(' ', '_')}_N{num_doc}_{grado_seccion.replace(' ', '_')}.docx"
                    st.session_state['ie_nombre_generado'] = ie_nombre
                    st.session_state['imagen_nanobanana'] = None
                    st.session_state['imagen_bytes'] = None
                    
                    st.success(f"✅ ¡{tipo_documento} generado con éxito!")

        except Exception as e:
            err_str = str(e)
            if "429" in err_str or "RESOURCE_EXHAUSTED" in err_str:
                st.warning("⏳ Límite de velocidad alcanzado. Por favor, espera 60 segundos y vuelve a intentarlo.")
            elif "404" in err_str or "NOT_FOUND" in err_str:
                st.error("⚠️ El modelo seleccionado no está disponible. Selecciona gemini-2.0-flash o gemini-2.5-flash.")
            else:
                st.error(f"❌ Ocurrió un error con la API de Google AI Studio: {err_str}")

# ==============================================================================
# DESPLIEGUE DE VISTA PREVIA Y DESCARGA PERMANENTE
# ==============================================================================
if st.session_state['resultado_md'] is not None:
    st.markdown("---")
    
    tab_preview, tab_download = st.tabs(["📄 Vista Previa (Permanente)", "📥 Descargar Afiche / Documento"])
    
    with tab_preview:
        # Muestra el afiche ilustrado generado por Nano Banana
        if st.session_state.get('imagen_nanobanana') is not None:
            st.markdown("### 🖼️ Afiche Educativo Ilustrado (Nano Banana AI)")
            st.image(st.session_state['imagen_nanobanana'], caption=f"Afiche para {grado_seccion} - {problema_contexto}", use_container_width=True)
            st.markdown("---")

        st.markdown(st.session_state['resultado_md'])
        
    with tab_download:
        # CASO A: Si el documento solicitado es el Afiche Nano Banana
        if st.session_state.get('tipo_doc_generado') == "Afiche Educativo de la Sesión (Nano Banana)":
            st.markdown("### 🖼️ Descarga tu Afiche Educativo Ilustrado")
            if st.session_state.get('imagen_bytes') is not None:
                st.download_button(
                    label="💾 Descargar Afiche Ilustrado en Alta Calidad (.jpg)",
                    data=st.session_state['imagen_bytes'],
                    file_name=f"Afiche_NanoBanana_{grado_seccion.replace(' ', '_')}.jpg",
                    mime="image/jpeg",
                    use_container_width=True
                )
                st.success("✨ ¡Tu afiche en JPG está listo para imprimir o enviar por WhatsApp a los alumnos!")
            else:
                st.warning("⚠️ No se pudo generar la foto del afiche. Por favor verifica los permisos de tu API Key de Google AI Studio.")

        # CASO B: Para los demás documentos en Word (Proyecto, Unidad, Sesión, Ficha)
        else:
            es_horizontal_doc = st.session_state['tipo_doc_generado'] in ["Proyecto de Aprendizaje", "Unidad de Aprendizaje (Modelo SARA)"]
            
            buffer_doc = markdown_to_docx(
                st.session_state['resultado_md'], 
                ie_nombre=st.session_state.get('ie_nombre_generado', ie_nombre),
                es_horizontal=es_horizontal_doc
            )
            
            st.download_button(
                label=f"💾 Descargar {st.session_state['tipo_doc_generado']} en Word (.docx)",
                data=buffer_doc,
                file_name=st.session_state['fname_clean'],
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
            st.info("💡 **Nota:** El documento Word generado incluye los recuadros y tablas en tonos pasteles.")
